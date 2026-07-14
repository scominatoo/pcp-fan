#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Gera checklist de homologação PCP em formato Word (.docx)."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill: str) -> None:
    """Aplica cor de fundo na célula (hex sem #)."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_info_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        set_cell_shading(table.rows[i].cells[0], 'E8ECF1')
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)


def add_checklist_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    col_widths_cm: list[float] | None = None,
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], '0F3460')
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.bold = True
                r.font.size = Pt(9)

    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val
            for p in table.rows[ri].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)


def add_result_line(doc: Document, label: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f'{label}: ').bold = True
    p.add_run('☐ Aprovado    ☐ Aprovado com ressalvas    ☐ Reprovado    ☐ Pendente')
    p.paragraph_format.space_after = Pt(6)


def add_signature_block(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, 'Aceite formal', level=1)
    add_paragraph(
        doc,
        'Declaro que participei da homologação descrita neste documento e que os '
        'resultados registrados refletem a validação realizada no ambiente de testes '
        '(pcp-homol), comparando com o sistema legado COBOL quando aplicável.',
    )
    doc.add_paragraph()

    sig_rows = [
        ('Nome do responsável FANANDRI', '_' * 50),
        ('Cargo / função', '_' * 50),
        ('Data', '_' * 20),
        ('Assinatura', '_' * 50),
        ('Nome do líder do projeto', '_' * 50),
        ('Assinatura', '_' * 50),
    ]
    add_info_table(doc, sig_rows)

    doc.add_paragraph()
    add_paragraph(
        doc,
        'Observação: o aceite pode ser formalizado por assinatura física, digital ou '
        'e-mail corporativo com cópia deste documento preenchido.',
        bold=False,
    )


def main() -> None:
    out_dir = Path(__file__).resolve().parent.parent / 'docs'
    out_file = out_dir / '17-checklist-homologacao-cliente.docx'

    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # Capa
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CHECKLIST DE HOMOLOGAÇÃO\n')
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(15, 52, 96)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run('Módulo PCP — Sistema Web (pcp-homol)\n')
    r2.font.size = Pt(14)
    r2.bold = True

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = sub2.add_run('Cliente: Indústria Metalúrgica FANANDRI Ltda.\n')
    r3.font.size = Pt(12)

    sub3 = doc.add_paragraph()
    sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = sub3.add_run(f'Versão do documento: {date.today().strftime("%d/%m/%Y")}')
    r4.font.size = Pt(11)
    r4.italic = True

    doc.add_paragraph()

    add_heading(doc, '1. Identificação da sessão', level=1)
    add_info_table(
        doc,
        [
            ('Data da homologação', ''),
            ('Local / ambiente', '☐ Presencial   ☐ Remoto'),
            ('URL do sistema novo', 'http://localhost:5175 (ajustar se necessário)'),
            ('Sistema legado (referência)', 'COBOL/DOS — FANANDRI'),
            ('Responsável FANANDRI (usuário-chave)', ''),
            ('Responsável técnico / líder do projeto', ''),
            ('Especialista legado (se houver)', ''),
        ],
    )

    add_heading(doc, '2. Instruções de uso', level=1)
    bullets = [
        'Marque ☐ Sim, ☐ Não ou ☐ Pendente em cada item conforme o teste.',
        'Para itens de amostra, anote códigos conferidos na coluna Observações.',
        'Compare sempre o sistema novo com o COBOL usando cópia dos dados migrados (.DAT).',
        'Não altere os mesmos registros nos dois sistemas ao mesmo tempo durante os testes.',
        'Registre divergências com detalhes (código, campo, valor legado vs. valor novo).',
        'Ao final de cada fase, marque o resultado: Aprovado / Ressalvas / Reprovado / Pendente.',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    add_heading(doc, '3. Critérios gerais de aceite', level=1)
    add_checklist_table(
        doc,
        ['#', 'Critério', 'Como verificar', 'OK', 'NOK', 'Pend.', 'Observações'],
        [
            ['G1', 'Contagem de registros', 'Totais migrados batem com o legado (±0)', '☐', '☐', '☐', ''],
            ['G2', 'Amostra campo a campo', '10 itens por módulo conferidos', '☐', '☐', '☐', ''],
            ['G3', 'Fluxo operacional', 'Usuário executa sem bloqueio indevido', '☐', '☐', '☐', ''],
            ['G4', 'Relatórios', 'Saída equivale ao impresso COBOL', '☐', '☐', '☐', ''],
        ],
        [1, 3.5, 4.5, 0.8, 0.8, 0.8, 3.5],
    )

    # FASE 0
    doc.add_page_break()
    add_heading(doc, 'Fase 0 — Fundação', level=1)
    add_paragraph(doc, 'Objetivo: confirmar que a base técnica e os dados migrados estão corretos.')
    add_checklist_table(
        doc,
        ['#', 'Item de validação', 'Referência legado', 'OK', 'NOK', 'Pend.', 'Observações'],
        [
            ['0.1', 'Banco PostgreSQL acessível e API respondendo', '—', '☐', '☐', '☐', ''],
            ['0.2', 'Total de produtos migrados: esperado 1.834', 'PCPA18I', '☐', '☐', '☐', ''],
            ['0.3', 'Total de matérias-primas migradas: esperado 3.298', 'PCPA22I', '☐', '☐', '☐', ''],
            ['0.4', 'Total de grupos: esperado 256', 'PCPA19I', '☐', '☐', '☐', ''],
            ['0.5', 'Total de classificações: esperado 35', 'PCPA20I', '☐', '☐', '☐', ''],
            ['0.6', 'Amostra de 10 produtos — campos principais', 'PC1018', '☐', '☐', '☐', ''],
            ['0.7', 'Amostra de 10 MPs — saldo, mín/máx, descrição', 'PC1022', '☐', '☐', '☐', ''],
        ],
        [1, 5.5, 2.5, 0.8, 0.8, 0.8, 3.5],
    )
    add_result_line(doc, 'Resultado Fase 0')

    # FASE 1
    add_heading(doc, 'Fase 1 — Cadastros', level=1)
    add_paragraph(doc, 'Telas: Produtos e Matéria-prima. Programas COBOL: PC1018, PC1022.')
    add_checklist_table(
        doc,
        ['#', 'Item de validação', 'Como testar no sistema novo', 'OK', 'NOK', 'Pend.', 'Observações'],
        [
            ['1.1', 'Listar produtos com busca e paginação', 'Menu Produtos', '☐', '☐', '☐', ''],
            ['1.2', 'Abrir detalhe/edição de produto', 'Clicar em um produto', '☐', '☐', '☐', ''],
            ['1.3', 'Código formatado (001-01-00001) correto', 'Comparar com COBOL', '☐', '☐', '☐', ''],
            ['1.4', 'Descrição, unidade, desenho cliente', 'Campo a campo', '☐', '☐', '☐', ''],
            ['1.5', 'Estoque mín/máx e quantidade', 'Campo a campo', '☐', '☐', '☐', ''],
            ['1.6', 'Listar matéria-prima com busca', 'Menu Matéria-prima', '☐', '☐', '☐', ''],
            ['1.7', 'Abrir detalhe/edição de MP', 'Clicar em uma MP', '☐', '☐', '☐', ''],
            ['1.8', 'Código MP (classe + item) correto', 'Comparar com COBOL', '☐', '☐', '☐', ''],
            ['1.9', 'Amostra 10 produtos — conferência completa', 'Planilha anexa', '☐', '☐', '☐', ''],
            ['1.10', 'Amostra 10 MPs — conferência completa', 'Planilha anexa', '☐', '☐', '☐', ''],
        ],
        [1, 4.5, 3.5, 0.8, 0.8, 0.8, 3.5],
    )
    add_paragraph(doc, 'Códigos da amostra — Produtos (preencher):', bold=True)
    add_info_table(
        doc,
        [(f'Produto {i}', '') for i in range(1, 11)],
    )
    add_paragraph(doc, 'Códigos da amostra — Matéria-prima (preencher):', bold=True)
    add_info_table(
        doc,
        [(f'MP {i}', '') for i in range(1, 11)],
    )
    add_result_line(doc, 'Resultado Fase 1')

    # FASE 2
    doc.add_page_break()
    add_heading(doc, 'Fase 2 — Processo produtivo e Ordens de Produção', level=1)
    add_paragraph(doc, 'Programas COBOL: PC1028 (cadastro OP), PC1041 (emissão). Total migrado: 72.000 OPs.')
    add_checklist_table(
        doc,
        ['#', 'Item de validação', 'Como testar', 'OK', 'NOK', 'Pend.', 'Observações'],
        [
            ['2.1', 'Listar OPs com filtros (aberta/baixada)', 'Menu Ordens de produção', '☐', '☐', '☐', ''],
            ['2.2', 'Abrir detalhe de OP migrada', 'Ex.: OP 1 ou 79814', '☐', '☐', '☐', ''],
            ['2.3', 'Operações da OP exibidas corretamente', 'Comparar com COBOL', '☐', '☐', '☐', ''],
            ['2.4', 'Cliente da OP correto', 'PCPA28II', '☐', '☐', '☐', ''],
            ['2.5', 'Criar nova OP por desenho do cliente', 'Nova OP — desenho 90531014', '☐', '☐', '☐', ''],
            ['2.6', 'Roteiro carregado automaticamente', 'Operações do processo', '☐', '☐', '☐', ''],
            ['2.7', 'Emitir OP (impressão/visualização)', 'Tela de emissão', '☐', '☐', '☐', ''],
            ['2.8', 'Bloqueio de emissão de OP já baixada', 'Tentar emitir OP baixada', '☐', '☐', '☐', ''],
            ['2.9', 'Amostra 10 OPs abertas — conferência', 'Código, produto, qtde, data', '☐', '☐', '☐', ''],
            ['2.10', 'Emissão impressa comparável ao PC1041', 'Imprimir e comparar', '☐', '☐', '☐', ''],
        ],
        [1, 4.5, 3.5, 0.8, 0.8, 0.8, 3.5],
    )
    add_paragraph(doc, 'Códigos OP da amostra (preencher):', bold=True)
    add_info_table(doc, [(f'OP {i}', '') for i in range(1, 11)])
    add_result_line(doc, 'Resultado Fase 2')

    # FASE 3
    add_heading(doc, 'Fase 3 — Baixas (OP e Matéria-prima)', level=1)
    add_paragraph(
        doc,
        'Programas COBOL: PC1132 (baixa operação), PC1028 baixa (consolidada), '
        'PC1076/PC1109 (baixa MP).',
    )
    add_heading(doc, '3a — Baixa de operações da OP', level=2)
    add_checklist_table(
        doc,
        ['#', 'Item de validação', 'Como testar', 'OK', 'NOK', 'Pend.', 'Observações'],
        [
            ['3.1', 'Visualizar baixas existentes (migradas)', 'Detalhe OP com baixas', '☐', '☐', '☐', ''],
            ['3.2', 'Registrar baixa de operação em OP aberta', 'POST baixa operação', '☐', '☐', '☐', ''],
            ['3.3', 'Encerrar OP (baixa consolidada)', 'Botão encerrar', '☐', '☐', '☐', ''],
            ['3.4', 'Bloqueio de nova baixa em OP encerrada', 'Tentar baixar novamente', '☐', '☐', '☐', ''],
            ['3.5', 'Peças produzidas na baixa consolidada', 'Comparar PCPA71I', '☐', '☐', '☐', ''],
            ['3.6', 'Amostra 5 OPs baixadas — dados consolidados', 'Campo a campo', '☐', '☐', '☐', ''],
        ],
        [1, 4.5, 3.5, 0.8, 0.8, 0.8, 3.5],
    )
    add_heading(doc, '3b — Baixa de matéria-prima', level=2)
    add_checklist_table(
        doc,
        ['#', 'Item de validação', 'Como testar', 'OK', 'NOK', 'Pend.', 'Observações'],
        [
            ['3.7', 'Listar MPs do roteiro na OP', 'Seção baixa MP', '☐', '☐', '☐', ''],
            ['3.8', 'Registrar baixa de MP', 'Debitar estoque', '☐', '☐', '☐', ''],
            ['3.9', 'Saldo de MP atualizado após baixa', 'Conferir cadastro MP', '☐', '☐', '☐', ''],
            ['3.10', 'Flag baixadaMp na OP', 'Detalhe da OP', '☐', '☐', '☐', ''],
            ['3.11', 'Amostra 5 MPs — saldo antes/depois', 'Movimento PCPA76I', '☐', '☐', '☐', ''],
        ],
        [1, 4.5, 3.5, 0.8, 0.8, 0.8, 3.5],
    )
    add_result_line(doc, 'Resultado Fase 3')

    # FASE 4
    doc.add_page_break()
    add_heading(doc, 'Fase 4 — Programação de entregas', level=1)
    add_paragraph(doc, 'Programas COBOL: PC1066, PC1133. Migrados: 1.383 registros (PCPA66I).')
    add_checklist_table(
        doc,
        ['#', 'Item de validação', 'Como testar', 'OK', 'NOK', 'Pend.', 'Observações'],
        [
            ['4.1', 'Listar programação com filtros de data', 'Menu Programação', '☐', '☐', '☐', ''],
            ['4.2', 'Resumo (programado, entregue, a produzir)', 'Painel superior', '☐', '☐', '☐', ''],
            ['4.3', 'Consultar atrasos', 'Aba Atrasos', '☐', '☐', '☐', ''],
            ['4.4', 'Criar nova programação', 'Nova programação', '☐', '☐', '☐', ''],
            ['4.5', 'Registrar entrega parcial/total', 'Detalhe + entrega', '☐', '☐', '☐', ''],
            ['4.6', 'Saldo pendente calculado corretamente', 'Comparar PC1066', '☐', '☐', '☐', ''],
            ['4.7', 'Amostra 10 programações — conferência', 'Data, produto, qtde', '☐', '☐', '☐', ''],
        ],
        [1, 4.5, 3.5, 0.8, 0.8, 0.8, 3.5],
    )
    add_result_line(doc, 'Resultado Fase 4')

    # FASE 5
    add_heading(doc, 'Fase 5 — Relatórios', level=1)
    add_paragraph(
        doc,
        'Programas COBOL: PC1078, PC1071, PC1135, PC1059, PC1067. '
        'Menu: Relatórios no sistema web.',
    )
    add_checklist_table(
        doc,
        ['#', 'Relatório', 'Programa COBOL', 'OK', 'NOK', 'Pend.', 'Observações'],
        [
            ['5.1', 'OP em aberto', 'PC1078', '☐', '☐', '☐', ''],
            ['5.2', 'OP baixadas', 'PC1071', '☐', '☐', '☐', ''],
            ['5.3', 'Produção por setor', 'PC1135', '☐', '☐', '☐', ''],
            ['5.4', 'Estoque MP crítico (mín/máx)', 'PC1059', '☐', '☐', '☐', ''],
            ['5.5', 'Programação sintética por mês', 'PC1067', '☐', '☐', '☐', ''],
            ['5.6', 'Totais do relatório batem com COBOL', 'Impresso legado', '☐', '☐', '☐', ''],
            ['5.7', 'Filtros por período funcionam', 'Data início/fim', '☐', '☐', '☐', ''],
            ['5.8', 'Impressão (botão Imprimir) legível', 'window.print', '☐', '☐', '☐', ''],
        ],
        [1, 4, 2.5, 0.8, 0.8, 0.8, 4],
    )
    add_result_line(doc, 'Resultado Fase 5')

    # FASE 6 + divergências
    add_heading(doc, 'Fase 6 — Aceite final', level=1)
    add_checklist_table(
        doc,
        ['#', 'Item', 'OK', 'NOK', 'Pend.', 'Observações'],
        [
            ['6.1', 'Todas as fases anteriores aprovadas ou com plano de correção', '☐', '☐', '☐', ''],
            ['6.2', 'Divergências críticas corrigidas', '☐', '☐', '☐', ''],
            ['6.3', 'Usuário-chave consegue operar o dia a dia sem apoio técnico', '☐', '☐', '☐', ''],
            ['6.4', 'Documentação básica de uso entregue', '☐', '☐', '☐', ''],
            ['6.5', 'Aceite formal assinado', '☐', '☐', '☐', ''],
        ],
        [1, 8, 0.8, 0.8, 0.8, 5],
    )
    add_result_line(doc, 'Resultado Fase 6 — ACEITE FINAL DO PCP')

    add_heading(doc, 'Registro de divergências', level=1)
    add_paragraph(
        doc,
        'Use esta tabela para anotar problemas encontrados durante a homologação.',
    )
    div_headers = ['#', 'Fase', 'Descrição da divergência', 'Gravidade', 'Status', 'Responsável', 'Prazo']
    div_rows = [[str(i), '', '', '☐ Alta ☐ Média ☐ Baixa', '☐ Aberta ☐ Corrigida', '', ''] for i in range(1, 11)]
    add_checklist_table(doc, div_headers, div_rows, [0.8, 1, 5, 2.5, 2, 2, 1.5])

    add_signature_block(doc)

    doc.save(out_file)
    print(f'Documento gerado: {out_file}')


if __name__ == '__main__':
    main()
