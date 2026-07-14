#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Gera manual do usuário PCP Homologação em Word (.docx)."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_title(doc: Document, text: str, size: int = 20) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(15, 52, 96)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_text(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style='List Number')


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run('Importante: ')
    run.bold = True
    run.font.color.rgb = RGBColor(180, 80, 0)
    p.add_run(text).font.size = Pt(10)


def add_table_2col(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    for i, (a, b) in enumerate(rows):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
        set_cell_shading(table.rows[i].cells[0], 'E8ECF1')
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)


def main() -> None:
    out_file = (
        Path(__file__).resolve().parent.parent
        / 'docs'
        / '18-manual-usuario-pcp-homol.docx'
    )

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # Capa
    add_title(doc, 'MANUAL DO USUÁRIO')
    add_title(doc, 'PCP Homologação — Sistema Web', size=14)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Indústria Metalúrgica FANANDRI Ltda.\n')
    r.font.size = Pt(12)
    r2 = p.add_run(f'Atualizado em {date.today().strftime("%d/%m/%Y")}')
    r2.italic = True
    r2.font.size = Pt(11)

    doc.add_page_break()

    # 1. Introdução
    add_heading(doc, '1. O que é este sistema?', level=1)
    add_text(
        doc,
        'O PCP Homologação é a versão web do módulo de Planejamento e Controle da '
        'Produção (PCP) da FANANDRI. Ele substitui, em ambiente de testes, as telas '
        'do sistema legado em COBOL, usando os mesmos dados migrados dos arquivos .DAT.',
    )
    add_text(doc, 'Com ele você pode:', bold=True)
    add_bullets(
        doc,
        [
            'Consultar e editar cadastros (produtos e matéria-prima)',
            'Criar, consultar e emitir ordens de produção (OP)',
            'Registrar baixas de produção e de matéria-prima',
            'Gerenciar programação de entregas',
            'Gerar relatórios equivalentes aos do COBOL',
        ],
    )
    add_note(
        doc,
        'Este é um ambiente de homologação (teste). O COBOL continua sendo a '
        'fonte de verdade em produção até o aceite formal do cliente.',
    )

    # 2. Ligar o sistema
    add_heading(doc, '2. Como ligar o sistema (antes de usar)', level=1)
    add_text(
        doc,
        'O sistema precisa de três partes rodando: banco de dados, API (servidor) e '
        'tela web (navegador). Siga a ordem abaixo.',
    )
    add_heading(doc, 'Passo 1 — Banco de dados (Docker)', level=2)
    add_numbered(
        doc,
        [
            'Abra o Docker Desktop e aguarde iniciar.',
            'Abra o Terminal do Mac.',
            'Execute os comandos:',
        ],
    )
    add_text(doc, 'cd /Users/scominato/pcp-homol/pcp-homol-api\ndocker compose up -d', bold=True)
    add_heading(doc, 'Passo 2 — API (servidor backend)', level=2)
    add_numbered(
        doc,
        [
            'Abra um novo Terminal (deixe o anterior aberto).',
            'Execute:',
        ],
    )
    add_text(
        doc,
        'cd /Users/scominato/pcp-homol/pcp-homol-api\n'
        'export $(grep -v \'^#\' .env | xargs)\n'
        'npm run start:dev',
        bold=True,
    )
    add_text(doc, 'Aguarde aparecer algo como: "Nest application successfully started".')
    add_heading(doc, 'Passo 3 — Tela web (frontend)', level=2)
    add_numbered(
        doc,
        [
            'Abra mais um Terminal.',
            'Execute:',
        ],
    )
    add_text(doc, 'cd /Users/scominato/pcp-homol/pcp-homol-web\nnpm run dev', bold=True)
    add_heading(doc, 'Passo 4 — Abrir no navegador', level=2)
    add_numbered(
        doc,
        [
            'Abra Chrome, Safari ou Edge.',
            'Acesse: http://localhost:5175',
            'Na página inicial, confira se API e Banco aparecem como "ok".',
        ],
    )
    add_table_2col(
        doc,
        [
            ('URL do sistema', 'http://localhost:5175'),
            ('API (técnico)', 'http://localhost:3001/api/health'),
            ('Problema comum', 'Se a tela não carregar, verifique se os 3 passos acima estão rodando'),
        ],
    )

    doc.add_page_break()

    # 3. Navegação
    add_heading(doc, '3. Navegação — menu principal', level=1)
    add_text(
        doc,
        'No topo de todas as telas há o menu com as áreas do PCP:',
    )
    add_table_2col(
        doc,
        [
            ('Produtos', 'Cadastro e consulta de peças acabadas'),
            ('Matéria-prima', 'Cadastro e consulta de MPs'),
            ('Ordens de produção', 'Ciclo completo da OP'),
            ('Programação', 'Planejamento e entregas'),
            ('Relatórios', 'Consultas e impressos'),
        ],
    )
    add_text(doc, 'Clique no nome da área ou em "PCP Homologação" para voltar à página inicial.')

    # 4. Produtos
    add_heading(doc, '4. Produtos', level=1)
    add_heading(doc, '4.1 Consultar produtos', level=2)
    add_numbered(
        doc,
        [
            'Clique em Produtos no menu.',
            'Use a caixa de busca para pesquisar por código, descrição ou desenho.',
            'Clique em Buscar.',
            'Use Anterior / Próxima para navegar entre páginas.',
        ],
    )
    add_heading(doc, '4.2 Ver ou editar um produto', level=2)
    add_numbered(
        doc,
        [
            'Na lista, clique no produto desejado.',
            'Confira código (formato 001-01-00001), descrição, unidade, estoques.',
            'Altere os campos necessários e salve (se disponível na tela).',
        ],
    )

    # 5. Matéria-prima
    add_heading(doc, '5. Matéria-prima', level=1)
    add_numbered(
        doc,
        [
            'Clique em Matéria-prima no menu.',
            'Busque pela descrição ou código da MP.',
            'Clique em um item para ver saldo, estoque mínimo/máximo e demais dados.',
            'Após baixa de MP em uma OP, o saldo aqui deve refletir o débito.',
        ],
    )

    doc.add_page_break()

    # 6. OP
    add_heading(doc, '6. Ordens de produção (OP)', level=1)
    add_text(
        doc,
        'A OP é o documento central da produção. No legado, os programas equivalentes '
        'são PC1028 (cadastro) e PC1041 (emissão).',
    )

    add_heading(doc, '6.1 Consultar OPs', level=2)
    add_numbered(
        doc,
        [
            'Menu → Ordens de produção.',
            'Opcional: marque "Somente abertas" para ver só OPs não baixadas.',
            'Digite código da OP, produto ou cliente e clique em Buscar.',
            'Clique no número da OP para abrir o detalhe.',
        ],
    )

    add_heading(doc, '6.2 Criar nova OP', level=2)
    add_numbered(
        doc,
        [
            'Em Ordens de produção, clique em Nova OP.',
            'O sistema sugere o próximo código de OP automaticamente.',
            'Informe o desenho do cliente (ex.: 90531014) no campo Produto.',
            'Clique em Buscar processo — o sistema carrega o roteiro de produção.',
            'Preencha quantidade, tipo (PRD, PRO, PIL, TRY) e nome do cliente.',
            'Se necessário, ajuste o equipamento de cada operação.',
            'Clique em Salvar OP — você será levado ao detalhe da OP criada.',
        ],
    )
    add_note(
        doc,
        'Se aparecer erro ao buscar processo, o produto pode não ter roteiro '
        'cadastrado no banco migrado.',
    )

    add_heading(doc, '6.3 Emitir (imprimir) a OP', level=2)
    add_numbered(
        doc,
        [
            'Abra o detalhe da OP.',
            'Clique em Emitir OP (ou link equivalente na tela).',
            'Revise dados: produto, quantidade, operações, equipamentos.',
            'Clique em Imprimir para gerar o documento (equivalente ao PC1041).',
            'OPs já baixadas não podem ser emitidas — o sistema bloqueia.',
        ],
    )

    add_heading(doc, '6.4 Baixar operações da OP', level=2)
    add_text(doc, 'Equivalente ao programa PC1132 do legado.')
    add_numbered(
        doc,
        [
            'No detalhe de uma OP aberta, vá à seção Baixas de operação.',
            'Selecione o número da operação.',
            'Informe a quantidade produzida (qtde saída).',
            'Clique em Registrar baixa.',
            'Repita para cada operação concluída.',
        ],
    )

    add_heading(doc, '6.5 Encerrar a OP (baixa consolidada)', level=2)
    add_text(doc, 'Equivalente à baixa final do PC1028 / registro PCPA71I.')
    add_numbered(
        doc,
        [
            'Ainda no detalhe da OP, vá à seção Encerrar OP.',
            'Informe peças produzidas, MP consumida e rolos (se aplicável).',
            'Marque se MP e produto acabado já foram baixados.',
            'Clique em Encerrar OP.',
            'Após encerrada, a OP não aceita novas baixas de operação.',
        ],
    )

    add_heading(doc, '6.6 Baixar matéria-prima na OP', level=2)
    add_text(doc, 'Equivalente ao PC1076 / PC1109.')
    add_numbered(
        doc,
        [
            'No detalhe da OP, vá à seção Baixa de matéria-prima.',
            'Selecione a MP (lista vem do roteiro do produto).',
            'Informe quantidade e rolos, se houver.',
            'Clique em Registrar baixa de MP.',
            'O estoque da MP no cadastro será debitado automaticamente.',
        ],
    )

    doc.add_page_break()

    # 7. Programação
    add_heading(doc, '7. Programação de entregas', level=1)
    add_text(doc, 'Equivalente aos programas PC1066 e PC1133.')
    add_heading(doc, '7.1 Consultar programação', level=2)
    add_numbered(
        doc,
        [
            'Menu → Programação.',
            'No topo, veja o resumo: programado, entregue, a produzir, saldo.',
            'Use busca e filtros de data; clique em Filtrar.',
            'Na aba Lista, navegue pelos registros.',
            'Na aba Atrasos, veja itens pendentes com data vencida.',
        ],
    )
    add_heading(doc, '7.2 Criar programação', level=2)
    add_numbered(
        doc,
        [
            'Clique em Nova programação.',
            'Informe data, quantidade e produto (ou desenho).',
            'Preencha plano, flag, pedido se necessário.',
            'Salve.',
        ],
    )
    add_heading(doc, '7.3 Registrar entrega', level=2)
    add_numbered(
        doc,
        [
            'Na lista, clique na programação desejada.',
            'Informe a quantidade entregue.',
            'Confirme — o sistema atualiza entregue, a produzir e saldo pendente.',
        ],
    )

    # 8. Relatórios
    add_heading(doc, '8. Relatórios', level=1)
    add_numbered(
        doc,
        [
            'Menu → Relatórios.',
            'Escolha o relatório desejado na lista.',
            'Preencha filtros (datas, código OP, seção, tipo de estoque).',
            'Clique em Gerar.',
            'Confira totais no topo e a tabela abaixo.',
            'Use Imprimir para saída em papel (comparar com COBOL na homologação).',
        ],
    )
    add_table_2col(
        doc,
        [
            ('OP em aberto', 'PC1078 — OPs não baixadas'),
            ('OP baixadas', 'PC1071 — OPs encerradas com dados de produção'),
            ('Produção por setor', 'PC1135 — quantidade por seção'),
            ('Estoque MP crítico', 'PC1059 — abaixo do mínimo ou acima do máximo'),
            ('Programação sintética', 'PC1067 — totais por mês'),
        ],
    )

    doc.add_page_break()

    # 9. Fluxo do dia a dia
    add_heading(doc, '9. Fluxo típico do dia a dia', level=1)
    add_text(doc, 'Sequência recomendada para uma produção completa:')
    add_numbered(
        doc,
        [
            'Consultar programação do período (o que precisa ser produzido).',
            'Criar ou localizar a OP do produto.',
            'Emitir a OP para a fábrica (impressão).',
            'À medida que as operações forem concluídas, registrar baixas de operação.',
            'Baixar matéria-prima consumida na OP.',
            'Encerrar a OP quando a produção terminar.',
            'Registrar entrega na programação, se aplicável.',
            'Consultar relatórios para conferência (OP abertas, produção por setor, etc.).',
        ],
    )

    # 10. Problemas comuns
    add_heading(doc, '10. Problemas comuns e soluções', level=1)
    add_table_2col(
        doc,
        [
            ('Tela mostra "Backend offline"', 'Suba a API (Passo 2 da seção 2)'),
            ('Erro de conexão / API', 'Verifique Docker e se a porta 3001 está livre'),
            ('Lista vazia', 'Confirme se a migração dos dados foi executada'),
            ('Não encontra processo na nova OP', 'Produto sem roteiro migrado (PCPA70XI)'),
            ('Não consigo emitir OP', 'OP pode estar baixada — só OP aberta emite'),
            ('Não consigo baixar operação', 'OP já encerrada ou operação inválida'),
            ('Saldo MP não mudou', 'Atualize a página; confira se a baixa foi confirmada'),
        ],
    )

    add_heading(doc, '11. Suporte e homologação', level=1)
    add_text(
        doc,
        'Durante a homologação, use o documento '
        '"17-checklist-homologacao-cliente.docx" para validar cada fase com a FANANDRI. '
        'Anote divergências e compare sempre com o sistema COBOL em paralelo, '
        'sem alterar os mesmos registros nos dois sistemas ao mesmo tempo.',
    )

    doc.save(out_file)
    print(f'Documento gerado: {out_file}')


if __name__ == '__main__':
    main()
